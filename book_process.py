import PyPDF2
import tempfile
from docx import Document
import re

# Função para extrair texto do PDF
def extract_text_from_pdf(pdf_file, start_page, end_page):
    text_list = []
    reader = PyPDF2.PdfReader(pdf_file)
    
    for page_number in range(start_page - 1, min(end_page, len(reader.pages))):
        page = reader.pages[page_number]
        text = page.extract_text()
        
        # Remover espaços extras no início e no fim de cada linha
        text = '\n'.join(line.strip() for line in text.splitlines())
        
        # Substituir quebras de linha únicas por espaços, preservando parágrafos
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
        
        # Normalizar quebras de parágrafo para exatamente duas quebras de linha
        text = re.sub(r'\n{2,}', '\n\n', text)
        
        # Remover hifenização de palavras no final das linhas
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        # Remover números de página isolados no final do texto
        text_lines = text.splitlines()
        while text_lines and text_lines[-1].strip().isdigit():
            text_lines.pop()
        text = '\n'.join(text_lines)
        
        # Substituir nomes específicos (pode incluir outros conforme necessário)
        replacements = {
            "Josélia Dantas": " ",
            "Nossa casa em Floriano": " ",
            "Infância e Adolescência com Meus Irmãos": " "
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Remover múltiplos espaços em branco
        text = re.sub(r' +', ' ', text)
        
        # Adicionar o texto processado à lista
        text_list.append(text)
    
    return text_list


# Função para aplicar marcações ao texto revisado
def apply_markup(text):
    # Exemplo de marcações: **negrito**, *itálico*, ~~tachado~~ e <mark>destacado</mark>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'~~(.*?)~~', r'<s>\1</s>', text)
    # Adicionar marcação com destaque
    text = re.sub(r'<mark>(.*?)</mark>', r'<highlight>\1</highlight>', text)
    return text

# Função para salvar a revisão como um arquivo Word com marcações
def save_as_word(revised_text_list, initial_page):
    doc = Document()
    doc.add_heading('Revisão do Livro', 0)

    for i, revised_text in enumerate(revised_text_list):
        doc.add_heading(f'Página {initial_page+i}', level=1)
        
        revised_text = apply_markup(revised_text)
        
        for paragraph_text in revised_text.split('\n\n'):
            paragraph = doc.add_paragraph()
            
            # Processar as marcações HTML
            parts = re.split(r'(</?[bis]>)|(</?highlight>)', paragraph_text)
            run = None
            highlight_active = False
            for part in parts:
                if part == '<b>':
                    run = paragraph.add_run()
                    run.bold = True
                elif part == '</b>':
                    run.bold = False
                elif part == '<i>':
                    run = paragraph.add_run()
                    run.italic = True
                elif part == '</i>':
                    run.italic = False
                elif part == '<s>':
                    run = paragraph.add_run()
                    run.strike = True
                elif part == '</s>':
                    run.strike = False
                elif part == '<highlight>':
                    highlight_active = True
                    run = paragraph.add_run()
                    run.font.highlight_color = 7  # Aplicar destaque (amarelo)
                elif part == '</highlight>':
                    highlight_active = False
                else:
                    run = paragraph.add_run(part)
                    if highlight_active:
                        run.font.highlight_color = 7  # Manter destaque ativo

    # Salvar o documento em um arquivo temporário
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name
